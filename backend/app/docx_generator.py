"""DOCX generation for material structures."""

import io
from typing import Any
from pydantic import BaseModel
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

from app.models import ExamStructure, DifferenzierungStructure


def generate_exam_docx(exam: ExamStructure) -> bytes:
    """Generate a DOCX file from an ExamStructure and return as bytes."""
    doc = Document()

    # Styles
    style = doc.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(11)

    # ── Title ──
    title = doc.add_heading(f"Klassenarbeit: {exam.thema}", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # ── Metadata ──
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run(f"Fach: {exam.fach}  |  Klasse: {exam.klasse}  |  Dauer: {exam.dauer_minuten} Minuten  |  Gesamtpunkte: {exam.gesamtpunkte}")
    meta.runs[0].font.size = Pt(10)
    meta.runs[0].font.color.rgb = RGBColor(100, 100, 100)

    # ── Hinweise ──
    if exam.hinweise:
        doc.add_heading("Hinweise", level=2)
        for h in exam.hinweise:
            doc.add_paragraph(h, style="List Bullet")

    doc.add_paragraph()  # spacer

    # ── Aufgaben ──
    doc.add_heading("Aufgaben", level=1)
    for i, task in enumerate(exam.aufgaben, 1):
        p = doc.add_paragraph()
        run = p.add_run(f"Aufgabe {i} ")
        run.bold = True
        run.font.size = Pt(12)

        badge = p.add_run(f"[AFB {task.afb_level}] ")
        badge.font.size = Pt(9)
        badge.font.color.rgb = RGBColor(80, 80, 80)

        pts = p.add_run(f"({task.punkte} Punkte)")
        pts.font.size = Pt(10)
        pts.font.color.rgb = RGBColor(120, 120, 120)

        doc.add_paragraph(task.beschreibung)
        doc.add_paragraph()  # spacer

    # ── Page break before Erwartungshorizont ──
    doc.add_page_break()

    # ── Erwartungshorizont ──
    doc.add_heading("Erwartungshorizont", level=1)

    table = doc.add_table(rows=1, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"

    hdr = table.rows[0].cells
    for cell, text in zip(hdr, ["Aufgabe", "AFB", "Erwartete Leistung", "Punkte"]):
        cell.text = text
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(10)

    for i, task in enumerate(exam.aufgaben, 1):
        row = table.add_row().cells
        row[0].text = str(i)
        row[1].text = task.afb_level
        row[2].text = "\n".join(f"• {e}" for e in task.erwartung)
        row[3].text = str(task.punkte)

    doc.add_paragraph()

    # ── Notenschlüssel ──
    doc.add_heading("Bewertungsraster / Notenschlüssel", level=2)

    note_table = doc.add_table(rows=1, cols=2)
    note_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    note_table.style = "Table Grid"

    hdr2 = note_table.rows[0].cells
    hdr2[0].text = "Note"
    hdr2[1].text = "Punkte"
    for cell in hdr2:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True

    for note, punkte in exam.notenschluessel.items():
        row = note_table.add_row().cells
        row[0].text = note
        row[1].text = punkte

    # ── Write to bytes ──
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def generate_diff_docx(diff: DifferenzierungStructure) -> bytes:
    """Generate a DOCX file from a DifferenzierungStructure."""
    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(11)

    # ── Title ──
    title = doc.add_heading(f"Differenziertes Material: {diff.thema}", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run(f"Fach: {diff.fach}  |  Klasse: {diff.klasse}")
    meta.runs[0].font.size = Pt(10)
    meta.runs[0].font.color.rgb = RGBColor(100, 100, 100)

    # ── Allgemeine Hinweise ──
    if diff.allgemeine_hinweise:
        doc.add_heading("Allgemeine Hinweise", level=2)
        for h in diff.allgemeine_hinweise:
            doc.add_paragraph(h, style="List Bullet")

    # ── Niveau colors ──
    niveau_colors = {
        "Basis": RGBColor(46, 125, 50),      # green
        "Mittel": RGBColor(245, 124, 0),      # orange
        "Erweitert": RGBColor(198, 40, 40),   # red
    }

    for niveau in diff.niveaus:
        doc.add_page_break()
        color = niveau_colors.get(niveau.niveau, RGBColor(0, 0, 0))

        heading = doc.add_heading(level=1)
        run = heading.add_run(f"Niveau: {niveau.niveau}")
        run.font.color.rgb = color

        info = doc.add_paragraph()
        info.add_run(f"Zeitaufwand: ca. {niveau.zeitaufwand_minuten} Minuten").font.size = Pt(10)

        if niveau.hinweise:
            for h in niveau.hinweise:
                doc.add_paragraph(h, style="List Bullet")

        doc.add_paragraph()

        for i, task in enumerate(niveau.aufgaben, 1):
            p = doc.add_paragraph()
            run = p.add_run(f"Aufgabe {i}: {task.aufgabe} ")
            run.bold = True
            run.font.size = Pt(12)

            pts = p.add_run(f"({task.punkte} Punkte)")
            pts.font.size = Pt(10)
            pts.font.color.rgb = RGBColor(120, 120, 120)

            doc.add_paragraph(task.beschreibung)

            if task.hilfestellung:
                hint = doc.add_paragraph()
                hint_run = hint.add_run(f"💡 Hilfestellung: {task.hilfestellung}")
                hint_run.font.size = Pt(10)
                hint_run.font.italic = True
                hint_run.font.color.rgb = RGBColor(100, 100, 100)

            doc.add_paragraph()

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def generate_generic_docx(structure: BaseModel, title: str = "Material") -> bytes:
    """Generate a DOCX from any Pydantic BaseModel structure.
    Renders fields as headings + content, lists as bullet points, nested models as sub-sections."""
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(11)

    heading = doc.add_heading(title, level=0)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

    _render_model(doc, structure, level=1)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# Field labels for German display
_LABELS = {
    "titel": "Titel", "thema": "Thema", "fach_thema": "Fach/Thema",
    "leitfrage": "Leitfrage", "hintergrund": "Hintergrund",
    "einfuehrung": "Einführung", "abschluss": "Abschluss",
    "lehrkraft_hinweise": "Hinweise für die Lehrkraft",
    "didaktische_hinweise": "Didaktische Hinweise",
    "lernziel": "Lernziel", "kerninhalt": "Kerninhalt",
    "theoretischer_hintergrund": "Theoretischer Hintergrund",
    "beobachtung": "Beobachtung", "auswertung": "Auswertung",
    "reflexion": "Reflexion", "handlungssituation": "Handlungssituation",
    "einstieg": "Einstieg", "loesung": "Lösung",
    "differenzierung": "Differenzierung", "weiterfuehrend": "Weiterführend",
    "spielname": "Spielname", "regeln": "Regeln", "inhalt": "Inhalt",
    "beschreibung": "Beschreibung", "hinweis": "Hinweis",
    "uebergang": "Übergang zum nächsten Rätsel",
    "hypothese": "Hypothese", "schwierigkeitsgrad": "Schwierigkeitsgrad",
    "offenheitsgrad": "Offenheitsgrad", "niveau": "Niveau",
    "beruf": "Beruf", "lernfeld": "Lernfeld", "zielgruppe": "Zielgruppe",
    "aufgabe": "Aufgabe", "kompetenzbereich": "Kompetenzbereich",
    "anforderungsniveau": "Anforderungsniveau", "sozialform": "Sozialform",
    "phase": "Phase", "medien": "Medien",
    "lehreraktivitaet": "Lehreraktivität", "schueleraktivitaet": "Schüleraktivität",
}

# Fields to skip (rendered as part of title or metadata)
_SKIP = {"titel", "nummer"}

# Fields that are simple metadata (rendered inline, not as heading)
_INLINE = {
    "thema", "fach_thema", "niveau", "schwierigkeitsgrad", "offenheitsgrad",
    "beruf", "lernfeld", "zielgruppe", "phase", "sozialform", "medien",
    "kompetenzbereich", "anforderungsniveau", "kategorie", "schwierigkeit",
}


def _render_model(doc: Document, model: BaseModel, level: int):
    """Recursively render a Pydantic model into a DOCX document."""
    data = model.model_dump()

    # Render inline fields as metadata paragraph
    inline_parts = []
    for key in _INLINE:
        if key in data and data[key]:
            label = _LABELS.get(key, key.replace("_", " ").title())
            inline_parts.append(f"{label}: {data[key]}")
    if inline_parts:
        meta = doc.add_paragraph(" | ".join(inline_parts))
        meta.runs[0].font.size = Pt(10) if meta.runs else None

    # Render remaining fields
    for key, val in data.items():
        if key in _SKIP or key in _INLINE or val is None:
            continue

        label = _LABELS.get(key, key.replace("_", " ").title())

        if isinstance(val, str):
            doc.add_heading(label, level=min(level, 4))
            doc.add_paragraph(val)
        elif isinstance(val, int):
            doc.add_paragraph(f"{label}: {val}")
        elif isinstance(val, list):
            if not val:
                continue
            if isinstance(val[0], str):
                doc.add_heading(label, level=min(level, 4))
                for item in val:
                    doc.add_paragraph(item, style="List Bullet")
            elif isinstance(val[0], dict):
                doc.add_heading(label, level=min(level, 4))
                for i, item in enumerate(val, 1):
                    # Check if it's a nested model with nummer
                    nr = item.get("nummer", i)
                    item_title = item.get("titel", item.get("aufgabe", item.get("phase", f"#{nr}")))
                    sub_heading = doc.add_heading(f"{nr}. {item_title}", level=min(level + 1, 5))
                    # Render each field of the nested item
                    for k, v in item.items():
                        if k in ("nummer", "titel") or v is None:
                            continue
                        sub_label = _LABELS.get(k, k.replace("_", " ").title())
                        if isinstance(v, str):
                            p = doc.add_paragraph()
                            p.add_run(f"{sub_label}: ").bold = True
                            p.add_run(v)
                        elif isinstance(v, int):
                            doc.add_paragraph(f"{sub_label}: {v}")
                        elif isinstance(v, list):
                            doc.add_paragraph(f"{sub_label}:")
                            for li in v:
                                doc.add_paragraph(str(li), style="List Bullet")
        elif isinstance(val, dict):
            doc.add_heading(label, level=min(level, 4))
            for k, v in val.items():
                doc.add_paragraph(f"{k}: {v}", style="List Bullet")


def generate_stundenplanung_docx(structure) -> bytes:
    """Special DOCX for Stundenplanung with Verlaufsplan table."""
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(11)

    heading = doc.add_heading(f"Stundenverlaufsplan: {structure.titel}", level=0)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run(f"Thema: {structure.fach_thema} | Zeitrahmen: {structure.zeitrahmen_minuten} Min.")

    doc.add_heading("Lernziel", level=2)
    doc.add_paragraph(structure.lernziel)

    # Verlaufsplan table
    doc.add_heading("Verlaufsplan", level=1)
    cols = ["Phase", "Zeit (Min)", "Lehreraktivität", "Schüleraktivität", "Sozialform", "Medien"]
    table = doc.add_table(rows=1, cols=len(cols))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, col_name in enumerate(cols):
        cell = table.rows[0].cells[i]
        cell.text = col_name
        for run in cell.paragraphs[0].runs:
            run.bold = True
            run.font.size = Pt(10)

    for phase in structure.phasen:
        row = table.add_row().cells
        row[0].text = phase.phase
        row[1].text = str(phase.zeit_minuten)
        row[2].text = phase.lehreraktivitaet
        row[3].text = phase.schueleraktivitaet
        row[4].text = phase.sozialform
        row[5].text = phase.medien

    if structure.didaktische_hinweise:
        doc.add_heading("Didaktische Hinweise", level=2)
        doc.add_paragraph(structure.didaktische_hinweise)

    if structure.materialien:
        doc.add_heading("Materialien", level=2)
        for m in structure.materialien:
            doc.add_paragraph(m, style="List Bullet")

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def generate_mystery_docx(structure) -> bytes:
    """Special DOCX for Mystery — prints cards as individual boxed items."""
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(11)

    heading = doc.add_heading(f"Mystery: {structure.titel}", level=0)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_heading("Leitfrage", level=1)
    p = doc.add_paragraph()
    run = p.add_run(structure.leitfrage)
    run.bold = True
    run.font.size = Pt(14)

    doc.add_heading("Hintergrund", level=2)
    doc.add_paragraph(structure.hintergrund)

    # Cards as table (one card per row, easy to cut out)
    doc.add_page_break()
    doc.add_heading("Informationskarten", level=1)
    doc.add_paragraph("(Zum Ausschneiden — jede Zeile ist eine Karte)")

    kategorie_colors = {
        "Fakt": RGBColor(46, 125, 50),
        "Hinweis": RGBColor(245, 124, 0),
        "Irreführung": RGBColor(198, 40, 40),
    }

    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for cell, text in zip(hdr, ["Nr.", "Kategorie", "Inhalt"]):
        cell.text = text
        for run in cell.paragraphs[0].runs:
            run.bold = True

    for karte in structure.karten:
        row = table.add_row().cells
        row[0].text = str(karte.nummer)
        row[1].text = karte.kategorie
        color = kategorie_colors.get(karte.kategorie, RGBColor(0, 0, 0))
        for run in row[1].paragraphs[0].runs:
            run.font.color.rgb = color
            run.bold = True
        row[2].text = karte.inhalt

    # Teacher page
    doc.add_page_break()
    doc.add_heading("Lösung (nur für Lehrkraft)", level=1)
    doc.add_paragraph(structure.loesung)

    doc.add_heading("Differenzierung", level=2)
    doc.add_paragraph(structure.differenzierung)

    doc.add_heading("Hinweise für die Lehrkraft", level=2)
    doc.add_paragraph(structure.lehrkraft_hinweise)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def generate_escape_room_docx(structure) -> bytes:
    """Special DOCX for Escape Room — narrative + sequential puzzles."""
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(11)

    heading = doc.add_heading(f"Escape Room: {structure.titel}", level=0)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run(f"Thema: {structure.thema} | Zeitrahmen: {structure.zeitrahmen_minuten} Min. | Schwierigkeit: {structure.schwierigkeitsgrad}")

    doc.add_heading("Einführung / Story", level=1)
    doc.add_paragraph(structure.einfuehrung)

    for raetsel in structure.raetsel:
        doc.add_heading(f"Rätsel {raetsel.nummer}: {raetsel.titel}", level=2)
        doc.add_paragraph(raetsel.beschreibung)

        if raetsel.material:
            hint = doc.add_paragraph()
            hint.add_run("Material: ").bold = True
            hint.add_run(raetsel.material)

    # Teacher page
    doc.add_page_break()
    doc.add_heading("Lösungen (nur für Lehrkraft)", level=1)

    for raetsel in structure.raetsel:
        doc.add_heading(f"Rätsel {raetsel.nummer}: {raetsel.titel}", level=2)
        p = doc.add_paragraph()
        p.add_run("Lösung: ").bold = True
        p.add_run(raetsel.loesung)
        p2 = doc.add_paragraph()
        p2.add_run("Hinweis (bei Bedarf): ").italic = True
        p2.add_run(raetsel.hinweis)
        if raetsel.uebergang:
            p3 = doc.add_paragraph()
            p3.add_run("Übergang: ").bold = True
            p3.add_run(raetsel.uebergang)

    doc.add_heading("Abschluss", level=1)
    doc.add_paragraph(structure.abschluss)

    doc.add_heading("Hinweise für die Lehrkraft", level=2)
    doc.add_paragraph(structure.lehrkraft_hinweise)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
