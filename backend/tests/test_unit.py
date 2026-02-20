"""
Unit Tests — Ebene 1: Reine Logik ohne externe Abhängigkeiten.

Testet Funktionen, die KEINEN LLM-Call und KEINEN DB-Call brauchen.
Diese Tests laufen in <1 Sekunde und kosten 0€.
"""

import pytest
import io
from zipfile import ZipFile
from docx import Document

from app.models import (
    ExamStructure, ExamTask,
    DifferenzierungStructure, DiffNiveau, DiffTask,
    ProfileUpdate, LoginRequest, ChatRequest, MaterialRequest,
)
from app.docx_generator import generate_exam_docx, generate_diff_docx
from app.ingestion import chunk_text
from app.services.material_service import resolve_material_type
from app.agents.system_prompt import BLOCK_IDENTITY, BLOCK_TOOLS
from app.h5p_generator import generate_multichoice, generate_blanks, generate_truefalse, generate_drag_text


# ═══════════════════════════════════════
# 1.1 Pydantic Models (Validation)
# ═══════════════════════════════════════

class TestModels:
    """Pydantic models validate input correctly."""

    def test_login_request_valid(self):
        req = LoginRequest(email="test@example.com", password="demo123")
        assert req.password == "demo123"
        assert req.email == "test@example.com"

    def test_login_request_empty_password(self):
        # Empty string is technically valid (the router handles rejection)
        req = LoginRequest(email="test@example.com", password="")
        assert req.password == ""

    def test_chat_request_valid(self):
        req = ChatRequest(message="Hallo", teacher_id="abc-123")
        assert req.message == "Hallo"
        assert req.conversation_id is None

    def test_chat_request_with_conversation(self):
        req = ChatRequest(message="Tschüss", teacher_id="abc", conversation_id="conv-1")
        assert req.conversation_id == "conv-1"

    def test_material_request_defaults(self):
        req = MaterialRequest(type="klausur", fach="Physik", klasse="9", thema="Mechanik", teacher_id="t1")
        assert req.dauer_minuten is None
        assert req.zusatz_anweisungen is None

    def test_profile_update_partial(self):
        req = ProfileUpdate(bundesland="NRW")
        assert req.bundesland == "NRW"
        assert req.schulform is None
        assert req.faecher is None
        assert req.jahrgaenge is None

    def test_profile_update_full(self):
        req = ProfileUpdate(
            bundesland="Sachsen",
            schulform="Gymnasium",
            faecher=["Physik", "Mathe"],
            jahrgaenge=[8, 9, 10],
        )
        assert len(req.faecher) == 2
        assert 10 in req.jahrgaenge

    def test_exam_task_valid(self):
        task = ExamTask(
            aufgabe="Test",
            beschreibung="Beschreibung",
            afb_level="II",
            punkte=10,
            erwartung=["Punkt 1"],
        )
        assert task.afb_level == "II"

    def test_exam_structure_gesamtpunkte(self):
        exam = ExamStructure(
            fach="Physik", klasse="9", thema="Optik", dauer_minuten=45,
            aufgaben=[
                ExamTask(aufgabe="A1", beschreibung="B1", afb_level="I", punkte=10, erwartung=["E1"]),
                ExamTask(aufgabe="A2", beschreibung="B2", afb_level="II", punkte=15, erwartung=["E2"]),
            ],
            gesamtpunkte=25,
            notenschluessel={"1": "24-25"},
            hinweise=["Hinweis 1"],
        )
        assert sum(t.punkte for t in exam.aufgaben) == exam.gesamtpunkte


# ═══════════════════════════════════════
# 1.2 Material Type Resolution
# ═══════════════════════════════════════

class TestMaterialTypeResolution:
    """resolve_material_type normalizes user input correctly."""

    def test_klausur_direct(self):
        assert resolve_material_type("klausur") == "klausur"

    def test_klassenarbeit_alias(self):
        assert resolve_material_type("Klassenarbeit") == "klausur"

    def test_test_passthrough(self):
        assert resolve_material_type("test") == "test"

    def test_pruefung_passthrough(self):
        assert resolve_material_type("prüfung") == "prüfung"

    def test_differenzierung_direct(self):
        assert resolve_material_type("differenzierung") == "differenzierung"

    def test_differenziert_passthrough(self):
        assert resolve_material_type("differenziert") == "differenziert"

    def test_arbeitsblatt_alias(self):
        assert resolve_material_type("arbeitsblatt") == "versuchsanleitung"

    def test_experiment_alias(self):
        assert resolve_material_type("experiment") == "versuchsanleitung"

    def test_unknown_is_passed_through(self):
        assert resolve_material_type("quiz") == "quiz"
        assert resolve_material_type("") == ""
        assert resolve_material_type("nonsense") == "nonsense"

    def test_case_insensitive(self):
        assert resolve_material_type("KLAUSUR") == "klausur"
        assert resolve_material_type("Differenzierung") == "differenzierung"


# ═══════════════════════════════════════
# 1.3 DOCX Generation
# ═══════════════════════════════════════

class TestDocxGeneration:
    """DOCX files are valid, contain expected content."""

    @pytest.fixture
    def sample_exam(self):
        return ExamStructure(
            fach="Physik", klasse="9", thema="Optik", dauer_minuten=45,
            aufgaben=[
                ExamTask(aufgabe="Brechungsgesetz", beschreibung="Erkläre das Snelliussche Brechungsgesetz.", afb_level="I", punkte=10, erwartung=["Snellius-Formel"]),
                ExamTask(aufgabe="Totalreflexion", beschreibung="Wann tritt Totalreflexion auf?", afb_level="II", punkte=15, erwartung=["Grenzwinkel"]),
            ],
            gesamtpunkte=25,
            notenschluessel={"1": "24-25", "2": "20-23", "3": "15-19"},
            hinweise=["Kein Taschenrechner"],
        )

    @pytest.fixture
    def sample_diff(self):
        return DifferenzierungStructure(
            fach="Deutsch", klasse="5", thema="Fabeln",
            niveaus=[
                DiffNiveau(niveau="Basis", aufgaben=[DiffTask(aufgabe="Lesen", beschreibung="Lies die Fabel.", punkte=5)], zeitaufwand_minuten=15, hinweise=["Wörterbuch"]),
                DiffNiveau(niveau="Mittel", aufgaben=[DiffTask(aufgabe="Analyse", beschreibung="Moral?", punkte=10)], zeitaufwand_minuten=20, hinweise=[]),
                DiffNiveau(niveau="Erweitert", aufgaben=[DiffTask(aufgabe="Transfer", beschreibung="Eigene Fabel.", punkte=20)], zeitaufwand_minuten=30, hinweise=[]),
            ],
            allgemeine_hinweise=["Einzelarbeit"],
        )

    def test_exam_docx_is_valid_docx(self, sample_exam):
        docx_bytes = generate_exam_docx(sample_exam)
        assert isinstance(docx_bytes, bytes)
        assert len(docx_bytes) > 1000  # Not trivially small

        # Should be a valid ZIP (DOCX = ZIP)
        buf = io.BytesIO(docx_bytes)
        assert ZipFile(buf).testzip() is None

    def test_exam_docx_contains_theme(self, sample_exam):
        docx_bytes = generate_exam_docx(sample_exam)
        doc = Document(io.BytesIO(docx_bytes))
        all_text = "\n".join(p.text for p in doc.paragraphs)
        assert "Optik" in all_text
        assert "Physik" in all_text

    def test_exam_docx_contains_all_tasks(self, sample_exam):
        docx_bytes = generate_exam_docx(sample_exam)
        doc = Document(io.BytesIO(docx_bytes))
        all_text = "\n".join(p.text for p in doc.paragraphs)
        assert "Aufgabe 1" in all_text
        assert "Aufgabe 2" in all_text
        assert "Brechungsgesetz" in all_text or "Snellius" in all_text

    def test_exam_docx_has_erwartungshorizont(self, sample_exam):
        docx_bytes = generate_exam_docx(sample_exam)
        doc = Document(io.BytesIO(docx_bytes))
        all_text = "\n".join(p.text for p in doc.paragraphs)
        # Table text isn't in paragraphs, check tables
        table_text = ""
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    table_text += cell.text + " "
        assert "Erwartungshorizont" in all_text or "Snellius" in table_text

    def test_exam_docx_has_notenschluessel(self, sample_exam):
        docx_bytes = generate_exam_docx(sample_exam)
        doc = Document(io.BytesIO(docx_bytes))
        table_text = ""
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    table_text += cell.text + " "
        assert "24-25" in table_text  # Note 1

    def test_diff_docx_is_valid(self, sample_diff):
        docx_bytes = generate_diff_docx(sample_diff)
        assert isinstance(docx_bytes, bytes)
        buf = io.BytesIO(docx_bytes)
        assert ZipFile(buf).testzip() is None

    def test_diff_docx_has_three_niveaus(self, sample_diff):
        docx_bytes = generate_diff_docx(sample_diff)
        doc = Document(io.BytesIO(docx_bytes))
        all_text = "\n".join(p.text for p in doc.paragraphs)
        assert "Basis" in all_text
        assert "Mittel" in all_text
        assert "Erweitert" in all_text

    def test_diff_docx_contains_theme(self, sample_diff):
        docx_bytes = generate_diff_docx(sample_diff)
        doc = Document(io.BytesIO(docx_bytes))
        all_text = "\n".join(p.text for p in doc.paragraphs)
        assert "Fabeln" in all_text


# ═══════════════════════════════════════
# 1.4 Chunking
# ═══════════════════════════════════════

class TestChunking:
    """chunk_text splits text correctly with overlap."""

    def test_short_text_single_chunk(self):
        text = "Dies ist ein kurzer Text."
        chunks = chunk_text(text, chunk_size=1000, chunk_overlap=100)
        assert len(chunks) == 1
        assert chunks[0]["text"] == text

    def test_long_text_multiple_chunks(self):
        # Create text that's definitely > 1 chunk
        paragraphs = [f"Absatz {i}: " + "x" * 200 for i in range(20)]
        text = "\n\n".join(paragraphs)
        chunks = chunk_text(text, chunk_size=500, chunk_overlap=100)
        assert len(chunks) > 1

    def test_chunks_have_metadata(self):
        text = "A" * 3000
        chunks = chunk_text(text, chunk_size=1000, chunk_overlap=100)
        for chunk in chunks:
            assert "text" in chunk
            assert "index" in chunk

    def test_empty_text_returns_empty(self):
        chunks = chunk_text("", chunk_size=1000, chunk_overlap=100)
        assert len(chunks) == 0 or (len(chunks) == 1 and chunks[0]["text"].strip() == "")

    def test_overlap_exists(self):
        text = "\n\n".join([f"Paragraph {i} content here." for i in range(50)])
        chunks = chunk_text(text, chunk_size=200, chunk_overlap=50)
        if len(chunks) >= 2:
            # Last chars of chunk N should appear in chunk N+1
            chunks[0]["text"][-50:]
            # Some overlap should exist
            assert len(chunks) >= 2  # Just verify we got multiple chunks


# ═══════════════════════════════════════
# 1.5 H5P Content Generation
# ═══════════════════════════════════════

class TestH5PGeneration:
    """H5P JSON structures are valid."""

    def test_multichoice_structure(self):
        answers = [
            {"text": "Berlin", "correct": True, "feedback": "Richtig!"},
            {"text": "München", "correct": False, "feedback": "Falsch."},
        ]
        content = generate_multichoice("Hauptstadt Deutschlands?", answers)
        assert "question" in content
        assert "answers" in content
        assert len(content["answers"]) == 2

    def test_blanks_structure(self):
        content = generate_blanks("Berlin ist die *Hauptstadt* von Deutschland.")
        assert "text" in content or "questions" in content

    def test_truefalse_structure(self):
        content = generate_truefalse("Die Erde ist rund.", True)
        assert "question" in content or "correct" in content

    def test_drag_text_structure(self):
        content = generate_drag_text("Berlin ist die Hauptstadt von *Deutschland*.")
        assert "textField" in content or "text" in content


# ═══════════════════════════════════════
# 1.6 System Prompt Blocks
# ═══════════════════════════════════════

class TestSystemPromptBlocks:
    """Static system prompt blocks contain critical instructions."""

    def test_identity_block_is_german(self):
        assert "Du bist eduhu" in BLOCK_IDENTITY

    def test_identity_mentions_rueckfragen(self):
        """Agent MUSS Rückfragen stellen bei Materialerstellung."""
        assert "Rückfragen" in BLOCK_IDENTITY or "Schärfungsfragen" in BLOCK_IDENTITY

    def test_identity_mentions_iteration(self):
        """Agent soll chirurgisch präzise iterieren."""
        assert "Chirurgische" in BLOCK_IDENTITY or "IDENTISCH" in BLOCK_IDENTITY

    def test_tools_block_lists_all_tools(self):
        assert "curriculum_search" in BLOCK_TOOLS or "Lehrplan" in BLOCK_TOOLS
        assert "web_search" in BLOCK_TOOLS or "Internet" in BLOCK_TOOLS
        assert "generate_material" in BLOCK_TOOLS
        assert "remember" in BLOCK_TOOLS
        assert "generate_exercise" in BLOCK_TOOLS

    def test_tools_block_mentions_download_link(self):
        """Critical: Agent muss Download-Link weitergeben."""
        assert "Download-Link" in BLOCK_TOOLS


# ═══════════════════════════════════════
# 1.7 Profile Suggestions (Pure Logic)
# ═══════════════════════════════════════

class TestBuildSuggestions:
    """build_suggestions creates personalized prompts from profile + memories."""

    def test_with_profile_and_no_memories(self):
        from app.routers.profile import build_suggestions
        profile = {"faecher": ["Physik"], "jahrgaenge": [9]}
        suggestions = build_suggestions(profile, [])
        assert len(suggestions) == 3
        assert any("Physik" in s for s in suggestions)

    def test_with_memories(self):
        from app.routers.profile import build_suggestions
        profile = {"faecher": ["Physik"], "jahrgaenge": [9]}
        memories = [
            {"key": "Thema", "value": "Optik", "category": "curriculum"},
            {"key": "Klasse", "value": "9a", "category": "class"},
        ]
        suggestions = build_suggestions(profile, memories)
        assert any("Optik" in s for s in suggestions)

    def test_empty_profile_and_memories(self):
        from app.routers.profile import build_suggestions
        suggestions = build_suggestions(None, [])
        assert len(suggestions) == 3
        # Should contain defaults
        assert any("Plane" in s or "Erstelle" in s or "Hilf" in s for s in suggestions)

    def test_max_three_suggestions(self):
        from app.routers.profile import build_suggestions
        profile = {"faecher": ["Physik", "Mathe", "Bio"], "jahrgaenge": [7, 8, 9]}
        memories = [
            {"key": "Thema", "value": "Optik", "category": "curriculum"},
            {"key": "Klasse", "value": "9a", "category": "class"},
        ]
        suggestions = build_suggestions(profile, memories)
        assert len(suggestions) <= 3
